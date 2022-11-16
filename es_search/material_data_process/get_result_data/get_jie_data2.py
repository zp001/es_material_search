# -*- coding: utf-8 -*-
import json
import os
from io import BytesIO
import datetime
import docx
import requests
from docx.document import Document
from docx.text.paragraph import Paragraph
from docx.image.image import Image
from docx.parts.image import ImagePart
from docx.oxml.shape import CT_Picture
from PIL import Image
import base64
import re
import tqdm

import material_data_process.get_result_data.comput_similarity as ed
from material_data_process.get_result_data import recog_small_title
from material_data_process.get_result_data.entity_label import get_position
from material_data_process.get_result_data.kmp import insert_label


def get_picture(document: Document, paragraph: Paragraph):
    """
    document 为文档对象
    paragraph 为内嵌图片的段落对象，比如第1段内
    """
    result_list = []
    img_list = paragraph._element.xpath('.//pic:pic')
    if len(img_list) == 0 or not img_list:
        return
    for i in range(len(img_list)):
        img: CT_Picture = img_list[i]
        embed = img.xpath('.//a:blip/@r:embed')[0]
        related_part: ImagePart = document.part.related_parts[embed]
        image: Image = related_part.image
        result_list.append(image)
    return result_list

def get_files(path):
    """
    读取word文档，对文档进行排序
    """
    file_name = []
    files = os.listdir(path)  # 采用listdir来读取所有文件
    for i in files:
        file_name.append(i)
    file_name.sort(key=lambda x: int(x.split('.')[0]))
    return file_name

def sort_files(files_list):
    """
    对前端上传的高拍仪word文档进行排序
        """
    n=len(files_list)
    if n%2==0:
        for i in range(n//2):
            files_list.insert(2*i+1,files_list[n//2+i])
            del files_list[n//2+i+1]
    if n%2==1:
        for i in range(int(n//2)):
            files_list.insert(2*i+1,files_list[int(n//2)+i+1])
            del files_list[int(n//2)+i+2]

    return files_list

def clean_text(para:str):
    """
        数据清洗
    """
    para=re.sub(r"\s|■|•",'',para)
    para = re.sub(r"\'", '', para)
    #para = para.encode('ascii', 'ignore').decode()  # 删除 unicode 字符（乱码
    #para =re.sub(r'[’!"#$%&\'()*+,-./:;<=>?@，。?★、…【】《》？“”‘’！\[\\\]^_`{|}~]+', ' ', para) # # 删除特殊字符
    #para = re.sub('\s{2,}', "", para) #  # 删除2个及以上的空格
    return para

def get_content(file):
    """
        读取word文档内容，以段落形式组织数据，每张图片等同一个段落。
        """
    data_list=[]
    for i in range(len(file.paragraphs)):
        para = file.paragraphs[i]
        image_list = get_picture(file, para)
        para=para.text.strip()
        para=para.replace(' ','')
        para=clean_text(para)
        p=process_picture(image_list)
        data_list.append(para)
        for pic in p:
            data_list.append(pic)
    return data_list

def process_flies(path,file_name):
    """
        得到文档地址
        """
    locations=[]
    for file in file_name:
        location = os.path.join(path, file)
        locations.append(location)
    return locations

def process_pig_parag(para):
    """
        清洗数据
        """
    para = para.replace(' ', '')
    para = clean_text(para)
    return para

def process_picture(image_list):
    """
        处理word文档图片
        """
    p=[]
    if image_list:
        for image in image_list:
            if image:
                blob = image.blob  # blob方法获取二进制流文件
                #Image.open(BytesIO(blob)).show()
                img_stream = base64.b64encode(blob)
                bs64 = "<img src=\"data:image/png;base64,"+img_stream.decode('utf-8')+"\">"
                p.append(bs64)
    return p

def get_data_list(locations):
    """
        处理上传的文档，得到分节标识'第'位置
        """
    all_data_list=[]
    all_data_list_index=[]
    for i in range(len(locations)):
        file = docx.Document(locations[i])
        data_list = get_content(file)
        data_list=[x for x in data_list if x!='']
        all_data_list.append(data_list)

    all_data_list=[i for j in all_data_list for i in j]
    for i in range(len(all_data_list)):
        if re.match(r'^[第][一二三四五六七八九十—]+节',all_data_list[i]):
            all_data_list_index.append(i)
    return all_data_list,all_data_list_index

def remove_pian(l):
    """
        去掉每章开头前面的描述
        """
    if re.match(r'^[第][一二三四五六七八九十—]+篇', l[0][0]):
        l=l[1:]
    return l

def get_data_patition(all_data_list,all_data_list_index):
    """
        将每章数据分成节，形成一个二维数组
        """
    new_data=[]
    for i in all_data_list_index:
        if all_data_list_index.index(i) ==0 and all_data_list_index[0]!=0:
            new_data.append(all_data_list[0:i])
        if all_data_list_index.index(i) != 0:
            new_data.append(all_data_list[all_data_list_index[all_data_list_index.index(i)-1]:i])
        if all_data_list_index.index(i) ==(len(all_data_list_index)-1):
            new_data.append(all_data_list[i:])
    new_data=remove_pian(new_data)

    return new_data

def get_material_catg():
    """
        得到物料名称和代码对应数据
        """
    file = '..//material_data_process//get_result_data//data//物料类别数据//物料类别.xls'
    #file = './/data//物料类别数据//物料类别.xls'
    bms_all_data, big_category_data, mid_category_data, small_category_data = ed.get_category_data(file)
    bms_data, bms_code = ed.get_data_dic(big_category_data, mid_category_data, small_category_data)
    return bms_data,bms_code

def get_small_title(con,section):
    """
        识别每节的一级标题，返回标题以及标题去掉引用后和节名称组合形成的组合标题，组合标题为后续搜索使用
        """
    title=[]
    combin_title=[]
    title_index = []
    for d in con:
        #if re.match(r'^[一二三四五六七八九十—]+[^\u4e00-\u9fa5]', d):
        if re.match(r'^[一二三四五六七八九十—]+、', d):
            index_d=con.index(d)
            title_index.append(index_d)
            s = d.split(d[1])[1]
            if '引用' in s:
                ind=s.index('引')
                new_s=s[:ind-1]
                conbin_s = section + new_s
                combin_title.append(conbin_s)
            if '引用' not in s:
                if '(' in s or '（' in s:
                    s.replace('(','（')
                    new_s=s.split('（')[0]
                    conbin_s = section + new_s
                    combin_title.append(conbin_s)
                else:
                    conbin_s=section+s
                    combin_title.append(conbin_s)
            title.append(d)
            i=con.index(d)
            new_d= "<h3>" + d + '</h3>'
            con[i]=new_d

    return title,con,combin_title,title_index

def title_smallTitle(title_index,small_title_index,title,small_title):
    ts=title_index+small_title_index
    ts=sorted(ts)
    ts_list=[]
    d=0
    for j in range(len(title_index)):
        ts_dic = {}
        if j!=len(title_index)-1:
            ind1=ts.index(title_index[j])
            ind2 = ts.index(title_index[j+1])
            s = len(ts[ind1 + 1:ind2])
            d = d + s
            ts_dic[title[j]] = small_title[d - s:d]
            #ts_dic[title[j]]=ts[ind1+1:ind2]
        else:
            ind1=ts.index(title_index[j])
            s = len(ts[ind1 + 1:])
            ts_dic[title[j]] = small_title[d - s:d]
            #ts_dic[title[j]] = ts[ind1 + 1:]
        ts_list.append(ts_dic)
    return ts_list

def all_title_tag(title_index,small_title_index,title,small_title):
    ts = title_index + small_title_index
    ts = sorted(ts)
    all_title = []
    all_title_tag=[]
    d = 0
    for j in range(len(title_index)):
        dt={}
        if j != len(title_index) - 1:
            ind1 = ts.index(title_index[j])
            ind2 = ts.index(title_index[j + 1])
            s = len(ts[ind1 + 1:ind2])
            d = d + s
            all_title.append(title[j])
            all_title.extend(small_title[d - s:d])
            dt['title']=title[j]
            dt['tag']='1'
            all_title_tag.append(dt)
            for st in small_title[d - s:d]:
                std={}
                std['title'] = st
                std['tag'] = '2'
                all_title_tag.append(std)
        else:
            ind1 = ts.index(title_index[j])
            s = len(ts[ind1 + 1:])
            all_title.append(title[j])
            all_title.extend(small_title[d - s:d])
            dt['title'] = title[j]
            dt['tag'] = '1'
            all_title_tag.append(dt)
            for st in small_title[d - s:d]:
                std = {}
                std['title'] = st
                std['tag'] = '2'
                all_title_tag.append(std)

    return all_title_tag,all_title


def get_para_pic(jie_list):
    """
        得到每节图片，并且图片与前面的文本段落对应:{text1：[p1,p2,...],text2:[p1,p2]}
        """
    all_pic = []
    para_pic = {}
    pic = []
    pic_ind = []
    for i in range(len(jie_list)):
        if jie_list[i].startswith('<img'):
            pic.append(jie_list[i])
            pic_ind.append(i)

        else:
            if len(pic_ind) != 0:
                para_pic[jie_list[pic_ind[0] - 1]] = pic
                all_pic.append(para_pic)
                para_pic = {}
                pic = []
                pic_ind = []
    return pic

def get_para_pic_list(jie_list):
    """
        得到整个节的图片，放到数组中
        """
    pic=[]
    for i in range(len(jie_list)):
        if jie_list[i].startswith('<img'):
            pic.append(jie_list[i])
    return pic

def remove_picture(s):
    """
        优化实体标签处理速度，将打标的节内容中图片代码去掉，形成新的不带图片代码的文本去打实体标签
        """
    s_list=s.split("|")
    pic_ind=[]
    pic=[]
    j = 0
    for i in range(len(s_list)):
        if s_list[i].startswith('<img'):
            pic.append(s_list[i])
            pic_ind.append(i)
        else:
            s_list[j] = s_list[i]
            j += 1
    new_s='|'.join(l for l in s_list[0:j])
    return pic,pic_ind,new_s


def insert_picture(pic,pic_ind,label_new_s):
    """
        将图片重新插入到打上标签的内容中
        """
    label_s_list = label_new_s.split("|")
    new_label_s_list = ['<p>' + w + '</p>' for w in label_s_list]
    for i in range(len(pic_ind)):
        new_label_s_list.insert(pic_ind[i],'<p>' +pic[i]+ '</p>')
    label_s = ''.join(l for l in new_label_s_list)
    return label_s

def get_quote(title):
    all_q = []
    for t in title:
        if '引用' in t:
            q = t.split('引用')[1][:-1]
            q=q[:-5]+'—'+q[-4:]
            if q not in all_q:
                all_q.append(q)
        else:
            if '(' in t or '（' in t:
                #if re.match(r'[(（][/a-zA-Z0-9][)）]', t):
                #p = re.compile(r'[（](.*)[）]', re.S) #匹配（）内内容，不加?,贪婪匹配,最终获取的为【最外层括号包含的内容 】
                p = re.compile(r'[（(](.*?)[）)]', re.S) #匹配（）内内容，加?，最小匹配
                t=re.findall(p,t)
                q=t[0]
                if q[-4:].isdigit():
                    q = q[:-5] + '—' + q[-4:]
                    if q not in all_q:
                        all_q.append(q)
            else:
                continue
    all_q=list(set(all_q))
    return all_q
def get_all_quote(jyy:str,title_quote,small_title_quote):
    for s in small_title_quote:
        if s in title_quote:
            small_title_quote.remove(s)
    title_quote.extend(small_title_quote)
    if jyy in title_quote:
        title_quote.remove(jyy)
        title_quote.insert(0,jyy)
    else:
        if jyy!='':
            title_quote.insert(0, jyy)
    title_quote=list(set(title_quote))
    return title_quote

def get_last_data(new_data):
    """
        得到处理后的每章数据[{},{}...]
        """
    zhang_list=[]
    bms_data, bms_code = get_material_catg()
    name_zhang=new_data[0][0].split('章')[1]
    for i in range(len(new_data)):
        jie_data={}
        jyy = ''
        if i==0:
            jie_data["chapter"] = name_zhang
            jie_data["section"] = name_zhang+'前言'
            #jie_data["content"] = new_data[i][1:]
            new_c = [p for p in new_data[i][1:] if not p.startswith('<img')]
            new_c_str = '|'.join(l for l in new_c)
            small_index_list = recog_small_title.rec_small_title(new_data[i][1:])
            s, small_title = recog_small_title.concat_cont(small_index_list, new_data[i][1:])
            pic,pic_ind,new_s=remove_picture(s)
            position, position_substr = get_position(new_s, name_zhang+'前言')
            label_new_s = insert_label(position, new_s)
            label_pic_s = insert_picture(pic, pic_ind, label_new_s)
            jie_data["content"] = label_pic_s
            jie_data["content_text"] = new_c_str
            jie_pic=get_para_pic_list(new_data[i][1:])
            jie_data["picture"] = jie_pic
            jie_data['material_name'] = []
            jie_data['material_code'] = []
            jie_data["quote"] = ''
            jie_data["title"] = []
            jie_data["title_quote"] = []
            jie_data["combin_title"] = []
            jie_data["small_title"] = []
            jie_data["small_title_quote"] = []
            jie_data["all_title"] = []
            jie_data["all_quote"] = []
            #jie_data["grade_all_title"] = []
            jie_data["all_title_tag"] = []
            now_time = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            jie_data["creatTime"] = now_time
        if i!=0:
            if '引用' in new_data[i][1] and not re.match(r'^[一二三四五六七八九十—]', new_data[i][1]):
                jyy = re.sub('[()（）]', '', new_data[i][1])
            else:
                jyy=''
            title,con,combin_title,title_index= get_small_title(new_data[i],new_data[i][0].split('节')[1])
            jie_data["chapter"] = name_zhang
            jie_data["section"]=new_data[i][0].split('节')[1]
            #jie_data["content"] = con[1:]
            new_c = [p for p in con[1:] if not p.startswith('<img')]
            new_c_str = '|'.join(l for l in new_c)
            small_index_list = recog_small_title.rec_small_title(con[1:])
            s, small_title = recog_small_title.concat_cont(small_index_list, con[1:])
            #ts_list=title_smallTitle(title_index, small_index_list, title, small_title)
            all_title_tag_list,all_title=all_title_tag(title_index, small_index_list, title, small_title)
            pic, pic_ind, new_s = remove_picture(s)
            position, position_substr = get_position(new_s, new_data[i][0].split('节')[1])
            label_new_s = insert_label(position, new_s)
            label_pic_s = insert_picture(pic, pic_ind, label_new_s)
            jie_data["content"] = label_pic_s
            jie_data["content_text"] = new_c_str
            jie_pic = get_para_pic_list(new_data[i][1:])
            jie_data["picture"] = jie_pic
            mat_sort_sim_list, cod_sort_sim_list = ed.match_all_data(new_data[i][0].split('节')[1], bms_data, bms_code)
            jie_data['material_name'] = mat_sort_sim_list
            jie_data['material_code'] = cod_sort_sim_list
            jie_data["quote"] = jyy
            jie_data["title"] = title
            title_quote=get_quote(title)
            jie_data["title_quote"] = title_quote
            jie_data["combin_title"] = combin_title
            jie_data["small_title"] = small_title
            small_title_quote = get_quote(small_title)
            jie_data["small_title_quote"] = small_title_quote
            all_quote=get_all_quote(jyy, title_quote, small_title_quote)
            #all_quote=list(set(all_quote))
            jie_data["all_quote"] = all_quote
            #jie_data["grade_all_title"] = ts_list
            jie_data["all_title_tag"] = all_title_tag_list
            jie_data["all_title"] = all_title
            now_time = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            jie_data["creatTime"] = now_time

        zhang_list.append(jie_data)

    return zhang_list

def get_result_data(zhang_list,l,processStatus):
    """
        添加一些前端传回的字段
        """
    new_zhang_list = []
    for t in zhang_list:
        updateTime = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        t["updateTime"] =updateTime
        t["creater"] = 'admin'
        t["modifiedName"] = 'admin'
        t["dataStatus"] = zhang_list[0]['dataStatus']
        t["versionNumber"] = l
        t["processStatus"] = processStatus
        t["desc"] = zhang_list[0]['desc']
        t["source"] = zhang_list[0]['source']
        t["top"] = ''
        t["unfreezReason"] = ''
        t["modifiedReason"] = ''
        new_zhang_list.append(t)

    return new_zhang_list

def load_data2es(article_list):
    """
        将数据写入到es
        """
    for i in tqdm.tqdm(article_list):
        data = json.dumps(i)
        try:
            base_url = "http://localhost:9200/" + "material" + "/" + "_doc" + "/"
            response = requests.post(base_url, headers={"Content-Type": "application/json"}, data=data.encode())
        except:
            print("失败")

if __name__ == "__main__":
    path="D://work//data//example_data//第一章//"
    file_name=get_files(path)
    locations=process_flies(path,file_name)
    all_data_list,all_data_list_index=get_data_list(locations)
    new_data=get_data_patition(all_data_list, all_data_list_index)
    zhang_list=get_last_data(new_data)
    #load_data2es(zhang_list)


'''
POST _analyze
{
  "analyzer": "ik_max_word",
  "text": "target="p"_blank"
}
'''
